import os
import datetime
import base64
from io import BytesIO
from pathlib import Path
from decimal import Decimal, InvalidOperation
import logging # Импортируем модуль логирования

from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required, user_passes_test
from django.contrib.auth.views import LoginView
from django.conf import settings
from django.urls import reverse_lazy
from django.http import HttpResponse, HttpResponseBadRequest, JsonResponse
from django.template.loader import render_to_string
from django.views.decorators.http import require_POST
from django.contrib import messages
from django.db.models import Max, Prefetch # Импортируем Prefetch
from django.utils.safestring import mark_safe # Для использования HTML в сообщениях

# Получаем логгер для приложения catalog
logger = logging.getLogger('catalog')


# === Импорты для работы с DOCX (python-docx) ===
from docx import Document
from docx.shared import Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
# Импорт для установки ширины колонок в DOCX
try:
    from docx.enum.table import WD_TABLE_COL_WIDTH
except ImportError:
    # Если константа недоступна (очень старая версия python-docx?)
    WD_TABLE_COL_WIDTH = None
    logger.warning("Could not import WD_TABLE_COL_WIDTH. Table width settings might not work.")

from docx.oxml.shared import OxmlElement, qn
# === Конец импортов DOCX ===

# === Импорты для работы с PDF (xhtml2pdf) ===
from xhtml2pdf import pisa
from xhtml2pdf.files import pisaFileObject
# === Конец импортов PDF ===

# Импорт для Pillow (оционально, для определения MIME-типа изображений)
try:
    from PIL import Image
    PIL_INSTALLED = True
except ImportError:
    PIL_INSTALLED = False
    logger.warning("Warning: Pillow not installed. Image MIME type detection will be based on file extension only.")

# === Убедитесь, что эти импорты есть и правильные из .models: ===
from .models import Category, Product, ProductModel, DocumentLog, get_next_document_number
# =====================================================


# =============================================================================
# ПАТЧ ДЛЯ РЕШЕНИЯ ПРОБЛЕМЫ С TTF НА WINDOWS (issue #623 в xhtml2pdf)
# =============================================================================
# Этот патч должен быть применен один раз при старте приложения.
# Проверяем флаг _patched_for_windows_ttf, чтобы избежать повторного применения
if not getattr(pisaFileObject, '_patched_for_windows_ttf', False):
    logger.info("Applying pisaFileObject patch for Windows TTF issue...")
    try:
        # Сохраняем оригинальную функцию на случай, если патч нужно отменить (хотя обычно не нужно)
        # original_getNamedFile = pisaFileObject.getNamedFile
        pisaFileObject.getNamedFile = lambda self: self.uri
        setattr(pisaFileObject, '_patched_for_windows_ttf', True) # Устанавливаем флаг
        logger.info("PisaFileObject patch applied successfully.")
    except Exception as e:
        logger.error(f"Error applying pisaFileObject patch: {e}")
else:
    logger.info("PisaFileObject patch already applied.")
# =============================================================================
# КОНЕЦ ПАТЧА
# =============================================================================


# Ваше пользовательское представление входа
class CustomLoginView(LoginView):
    template_name = 'registration/login.html'

    def get_redirect_url(self):
        # Перенаправляем суперпользователей в админку, остальных - на каталог
        redirect_to = self.request.GET.get(self.redirect_field_name, '')
        if not redirect_to:
             if self.request.user.is_superuser:
                 return reverse_lazy('admin:index')
             return reverse_lazy('catalog_selection')
        return redirect_to


# Ваше представление для домашней страницы
@login_required
def home_view(request):
    """Отображает домашнюю страницу после входа."""
    # Поскольку LOGIN_REDIRECT_URL настроен на catalog_selection, эта вьюха может
    # быть достигнута только если пользователь явно перешел по /home/ или
    # если LOGIN_REDIRECT_URL изменен. Можно просто редиректить на каталог.
    return redirect('catalog_selection')


# Выделенная функция для сбора данных выбранных позиций и расчета сумм
# Используется как в catalog_selection_view, так и в render_selection_summary_partial
# и функциях генерации документов (для получения model, quantity, price, total).
# Она НЕ будет получать image_path для generate_docx_offer в этой версии.
def get_selected_items_and_total(selection_data_copy):
    """
    Считывает данные из словаря selection_data_copy (копии сессии),
    получает модели из БД, валидирует данные, рассчитывает суммы
    и возвращает список словарей выбранных элементов, общую сумму и список сообщений.
    Функция работает с КОПИЕЙ словаря сессии и возвращает измененную копию,
    если были удалены некорректные элементы (например, элементы с quantity <= 0
    или с ID, которых нет в БД).
    """
    items = []
    total = Decimal('0.00')
    messages_list = [] # Собираем сообщения (ошибки валидации данных в сессии)

    # Получаем ID моделей из ключей selection_data_copy, которые выглядят как числа
    model_ids_in_selection = [int(model_id_str) for model_id_str in selection_data_copy.keys() if str(model_id_str).isdigit()]

    # Если выбор пуст, сразу возвращаем
    if not model_ids_in_selection:
         return items, total, messages_list, selection_data_copy


    # Получаем все выбранные модели одним запросом для эффективности
    # Используем select_related и prefetch_related для связанных объектов
    selected_models_queryset = ProductModel.objects.select_related('product').prefetch_related('product__category').filter(id__in=model_ids_in_selection)
    selected_models_dict = {str(model.id): model for model in selected_models_queryset} # Ключи - строки для соответствия сессии


    # Итерируем по КОПИИ selection_data_copy, чтобы можно было изменять ее
    # внутри цикла (например, удалять некорректные элементы)
    for model_id_str, item_data in list(selection_data_copy.items()):
         try:
            # Проверяем, что item_data является словарем (для защиты от старой структуры или поврежденных данных)
            if not isinstance(item_data, dict):
                logger.warning(f"Skipping invalid item_data format for ID {model_id_str} in session: {item_data}")
                # Удаляем некорректный элемент из сессии копии
                if model_id_str in selection_data_copy:
                    del selection_data_copy[model_id_str]
                    messages_list.append({'tags': 'warning', 'message': f"Предупреждение: Удалена некорректная позиция (ID: {model_id_str})."})
                continue # Пропускаем


            model = selected_models_dict.get(model_id_str)

            if not model:
                 logger.warning(f"Model ID {model_id_str} not found in DB during selection processing. Removing from session copy.")
                 if model_id_str in selection_data_copy:
                     del selection_data_copy[model_id_str]
                     messages_list.append({'tags': 'warning', 'message': f"Предупреждение: Позиция с ID {model_id_str} не найдена и удалена из выбора."})
                 continue

            # Читаем quantity и price из item_data словаря
            quantity = item_data.get('quantity', 0)
            price_str = item_data.get('price') # Может быть None или пустой строкой

            try:
                 # --- Валидация и приведение количества ---
                 # Обрабатываем None, нечисло, пустую строку. Убедимся, что количество - неотрицательное целое.
                 quantity = int(quantity) if quantity is not None and str(quantity).strip().isdigit() else 0
                 if quantity < 0: quantity = 0 # Негативное количество приравниваем к 0

                 # --- Валидация и приведение цены ---
                 unit_price = model.price # Цена по умолчанию из БД, если в сессии нет или она некорректна
                 price_source = "DB" # Источник цены для логирования

                 if price_str is not None and price_str != '': # Если в сессии есть ключ 'price' и он не пустая строка
                     try:
                          price_from_session = Decimal(price_str)
                          if price_from_session >= 0:
                               unit_price = price_from_session # Используем цену из сессии, если она валидна и >= 0
                               price_source = "Session"
                          else:
                               logger.warning(f"Negative price ({price_from_session}) found in session for model ID {model_id_str} ('{model.name}'). Using DB price.")
                               messages_list.append({'tags': 'warning', 'message': f"Обнаружена некорректная цена для позиции '{model.name}'. Использована цена из каталога."})
                               price_source = "DB (Invalid Session Price)"
                     except (ValueError, InvalidOperation):
                          logger.warning(f"Invalid price format '{price_str}' in session for model ID {model_id_str} ('{model.name}'). Using DB price.")
                          messages_list.append({'tags': 'warning', 'message': f"Предупреждение: некорректная цена для позиции '{model.name}' в выборе. Использована цена из каталога."})
                          price_source = "DB (Invalid Session Price Format)"
                 elif price_str is not None and price_str == '':
                      # Пустая строка цены в сессии означает "использовать цену из БД"
                      unit_price = model.price
                      price_source = "DB (Empty Session Price)"
                 # Если price_str is None, также используем цену из БД (поведение по умолчанию)


            except (ValueError, InvalidOperation): # Общая ошибка приведения типов
                 logger.warning(f"Error converting quantity ({quantity}) or price ({price_str}) for model ID {model_id_str} ('{model.name}'). Skipping item.")
                 messages_list.append({'tags': 'warning', 'message': f"Ошибка обработки количества/цены для позиции '{model.name}'. Позиция пропущена."})
                 # Удаляем некорректный элемент из сессии копии
                 if model_id_str in selection_data_copy:
                      del selection_data_copy[model_id_str]
                 continue


            if quantity <= 0:
                 # Если количество <= 0, удаляем элемент из сессии, если он там есть
                 if model_id_str in selection_data_copy:
                     del selection_data_copy[model_id_str]
                     logger.warning(f"Removing item ID {model_id_str} ('{model.name}') with quantity <= 0 from session copy.")
                 continue # Пропускаем элементы с нулевым количеством

            # Все корректно, добавляем элемент в список для отображения/генерации
            item_total = unit_price * quantity
            items.append({
                'model': model,
                'quantity': quantity,
                'unit_price': unit_price, # Это эффективная цена (из сессии или БД)
                'item_total': item_total,
                # 'image_path': image_path, # <--- НЕ ДОБАВЛЯЕМ image_path здесь
            })
            total += item_total
            # logger.debug(f"Item processed: {model.name}, Qty: {quantity}, Unit Price: {unit_price} ({price_source}), Item Total: {item_total}")


         except ValueError: # Ловим ошибки при int(model_id_str)
             logger.warning(f"Invalid model ID key in session: {model_id_str}. Removing from session copy.")
             if model_id_str in selection_data_copy:
                  del selection_data_copy[model_id_str]
             messages_list.append({'tags': 'warning', 'message': f"Предупреждение: Удалена некорректная позиция из выбора."})
             continue
         except Exception as e:
              logger.error(f"An unexpected error occurred while processing selected item ID {model_id_str}: {e}", exc_info=True) # Добавляем exc_info=True для вывода стека
              if model_id_str in selection_data_copy:
                  del selection_data_copy[model_id_str]
              messages_list.append({'tags': 'error', 'message': f"Произошла неожиданная ошибка при обработке позиции (ID: {model_id_str}): {e}. Удалена из выбора."})
              continue

    # Сортируем выбранные элементы для единообразия отображения/документа
    items = sorted(items, key=lambda x: (x['model'].product.category.name, x['model'].product.name, x['model'].name))

    # Важно: возвращаем измененную копию сессии, т.к. в ней могли быть удалены некорректные элементы
    return items, total, messages_list, selection_data_copy


# --- Ваше объединенное представление для каталога и выбора ---
@login_required
def catalog_selection_view(request):
    """
    Отображает список всех моделей продуктов, сгруппированных по категориям и продуктам,
    и текущий выбор пользователя с возможностью редактирования цен.
    """
    logger.info("Executing catalog_selection_view...")
    # Используем select_related и prefetch_related для оптимизации запросов
    product_models = ProductModel.objects.all().select_related('product').prefetch_related('product__category').order_by('product__category__name', 'product__name', 'name')
    logger.info(f"Query returned {product_models.count()} ProductModel objects.")

    structured_catalog = {}
    for model in product_models:
        # Убедимся, что есть связанные Product и Category
        if model.product and model.product.category:
             category = model.product.category
             product = model.product
             if category not in structured_catalog:
                 structured_catalog[category] = {}
             if product not in structured_catalog[category]:
                 structured_catalog[category][product] = []
             structured_catalog[category][product].append(model)
        else:
             logger.warning(f"Warning: Model ID {model.id} does not have a linked product or category. Skipping in catalog view.")


    logger.info(f"Structured catalog has {len(structured_catalog)} categories.")

    # === Обработка сессии ===
    selection = request.session.get('selection', {})

    # Для совместимости с старой структурой сессии, преобразуем ее при загрузке, если нужно
    # Проверяем, если selection не пуст и первый элемент имеет старую структуру (quantity - int)
    # (Это может быть упрощено или удалено после полного перехода)
    if selection and isinstance(list(selection.values())[0], (int, str)):
         logger.info("Detected old selection structure. Converting...")
         new_selection = {}
         old_selection_keys = list(selection.keys()) # Получаем ключи до начала модификации
         # Фильтруем ключи, которые могут быть приведены к int
         valid_old_keys = [k for k in old_selection_keys if str(k).isdigit()]
         if valid_old_keys:
             old_selection_models = ProductModel.objects.filter(id__in=[int(k) for k in valid_old_keys])
             old_models_dict = {str(model.id): model for model in old_selection_models}
         else:
             old_models_dict = {} # Если нет валидных ключей, словарь пуст

         is_conversion_modified = False # Флаг, чтобы знать, изменилась ли сессия во время конвертации

         for model_id_str, quantity in list(selection.items()): # Итерируем по копии, чтобы можно было удалять из оригинала
             try:
                 model = old_models_dict.get(model_id_str)
                 if model:
                    # Попытка преобразовать количество, если оно число или строка-число
                    q = int(quantity) if isinstance(quantity, (int, str)) and str(quantity).strip().isdigit() else 0
                    if q > 0: # Добавляем только если количество положительное
                        # При конвертации цена берется из БД
                        new_selection[model_id_str] = {'quantity': q, 'price': str(model.price)}
                    else:
                        # Удаляем позиции с некорректным или нулевым количеством из старой сессии при конвертации
                         if model_id_str in selection: # Удаляем из оригинальной сессии
                             del selection[model_id_str]
                             is_conversion_modified = True
                             logger.warning(f"Warning: Removed old selection item ID {model_id_str} with non-positive quantity ({quantity}) during conversion.")

                 else:
                     # Если модель не найдена в БД, удаляем ее из старой сессии
                     if model_id_str in selection: # Удаляем из оригинальной сессии
                         del selection[model_id_str]
                         is_conversion_modified = True
                         logger.warning(f"Warning: Old selection item ID {model_id_str} not found in DB during conversion. Removed.")
                         messages.warning(request, f"Предупреждение: Позиция с ID {model_id_str} не найдена и удалена из выбора.")
                         continue # Пропускаем
             except ValueError: # Ловим ошибки при int(model_id_str) или int(quantity) если проверки выше пропущены
                 if model_id_str in selection: # Удаляем из оригинальной сессии
                     del selection[model_id_str]
                     is_conversion_modified = True
                     logger.warning(f"Warning: Old selection item ID {model_id_str} has non-integer key or quantity ({quantity}). Removed during conversion.")
                     messages.warning(request, f"Предупреждение: Удалена некорректная позиция из выбора.")
                 continue
             except Exception as e:
                  logger.error(f"An unexpected error occurred during old selection conversion for ID {model_id_str}: {e}", exc_info=True)
                  if model_id_str in selection: # Удаляем из оригинальной сессии
                      del selection[model_id_str]
                      is_conversion_modified = True
                      messages.error(request, f"Произошла неожиданная ошибка при обработке позиции (ID: {model_id_str}): {e}. Удалена из выбора.")
                  continue # Пропускаем

         # Перезаписываем сессию новой структурой
         request.session['selection'] = new_selection # Перезаписываем сессию сконвертированными данными
         selection = new_selection # Обновляем локальную переменную
         if is_conversion_modified:
              request.session.modified = True
              logger.info("Old selection structure converted and session modified due to cleanup.")
         else:
              # Сессия могла измениться (например, только структура поменялась), но не помечена modified
              # Лучше пометить всегда, если конверсия произошла.
              request.session.modified = True # Всегда помечаем, если конверсия произошла
              logger.info("Old selection structure converted.")

    # Конец преобразования старой структуры

    # Теперь обрабатываем (или повторно обрабатываем после конвертации) сессию в новом формате
    # Используем get_selected_items_and_total для сбора данных и их форматирования
    # Передаем копию сессии, чтобы get_selected_items_and_total могла ее модифицировать (удалять некорректные)
    selected_items, total_price, messages_from_processing, updated_selection_copy = get_selected_items_and_total(selection.copy())

    # Если в процессе get_selected_items_and_total были удалены элементы из сессии копии,
    # обновляем оригинальную сессию и помечаем ее как измененную
    if updated_selection_copy != selection:
         request.session['selection'] = updated_selection_copy
         request.session.modified = True
         logger.info("Selection modified during catalog view load (cleaned up invalid items via get_selected_items_and_total).")
         # Важно: Обновляем локальную переменную selection тоже, чтобы шаблон получил актуальные данные
         selection = updated_selection_copy


    # Добавляем собранные сообщения (об ошибках валидации в сессии) в стандартные Django messages
    # Они будут отображены на странице при обычной загрузке.
    for msg_data in messages_from_processing:
         level = messages.INFO # По умолчанию
         if msg_data.get('tags') == 'warning': level = messages.WARNING
         elif msg_data.get('tags') == 'error': level = messages.ERROR
         messages.add_message(request, level, msg_data['message'])


    logger.info(f"Selected items list for summary has {len(selected_items)} items.")
    logger.info(f"Calculated total price: {total_price}")

    context = {
        'structured_catalog': structured_catalog,
        'selection': selection, # Передаем актуальную структуру сессии (возможно, обновленную после чистки)
        'selected_items': selected_items,
        'total_price': total_price,
        # ajax_messages больше не передаем в основной шаблон, т.к. используем стандартные messages
    }
    logger.info("Rendering catalog_selection.html...")
    return render(request, 'catalog/catalog_selection.html', context)


@require_POST
@login_required
def update_selection_view(request):
    """
    Обновляет количество или цены выбранных моделей в сессии по AJAX запросу
    и возвращает обновленный HTML-фрагмент сводки.
    Принимает POST параметры:
    - action: 'add', 'remove', 'set', 'remove_all', 'clear', 'apply_prices'
    - model_id: (для кол-ва) ID модели
    - {model_id_str}: {price_str} (для apply_prices) Пары ID модели и цены
    """
    logger.info(f"\nExecuting AJAX update_selection_view for user {request.user.username}")
    logger.debug(f"Received POST data: {request.POST}")

    action = request.POST.get('action')

    # Получаем текущую сессию
    selection = request.session.get('selection', {})
    logger.debug(f"Current selection before update: {selection}")

    # Хелпер для рендеринга сводки после операции
    # Собирает данные с помощью get_selected_items_and_total и рендерит частичный шаблон
    # Принимает список сообщений для отображения в AJAX ответе
    def render_and_return_partial(msgs_list=None):
         if msgs_list is None: msgs_list = []
         # Сбор данных и расчет сумм происходит внутри get_selected_items_and_total
         # Передаем КОПИЮ сессии, чтобы get_selected_items_and_total могла ее модифицировать (удалять некорректные)
         selected_items, total_price, messages_from_processing, updated_selection_copy = get_selected_items_and_total(selection.copy())

         # Если в процессе get_selected_items_and_total были удалены элементы из сессии копии,
         # обновляем оригинальную сессию и помечаем ее как измененную
         if updated_selection_copy != selection:
              request.session['selection'] = updated_selection_copy
              request.session.modified = True
              logger.info("Selection modified during partial render (cleaned up invalid items).")
              # Важно: Обновляем локальную переменную selection тоже, чтобы частичный шаблон получил актуальные данные
              #selection = updated_selection_copy # Нет, здесь не нужно, частичный шаблон рендерится с updated_selection_copy


         # Добавляем сообщения, собранные во время обработки, в список сообщений для AJAX ответа
         messages_list.extend(messages_from_processing)

         context = {
             'selected_items': selected_items,
             'total_price': total_price,
             'selection': updated_selection_copy, # Передаем актуальную структуру сессии из get_selected_items_and_total
             'ajax_messages': messages_list, # Передаем сообщения для отображения в AJAX ответе
         }

         html_content = render_to_string('catalog/selection_summary_partial.html', context, request=request)
         return HttpResponse(html_content, status=200)


    # === Обработка действия 'clear' ===
    if action == 'clear':
        request.session['selection'] = {}
        request.session.modified = True
        logger.info("Action 'clear': Selection has been emptied.")
        # Не добавляем стандартное сообщение здесь, используем только ajax_messages
        # messages.success(request, "Выбор очищен.") # Стандартное сообщение
        return render_and_return_partial([{'tags': 'success', 'message': "Выбор очищен."}])

    # === Обработка действия 'apply_prices' ===
    if action == 'apply_prices':
        logger.info("Action 'apply_prices' received.")
        # Проходим по всем POST параметрам. Ключи, которые являются цифрами, считаем ID моделей.
        # URLSearchParams отправляет данные как key=value, где key - это model_id_str
        prices_from_request = {}
        # Итерируем по request.POST.items(), чтобы получить все пары ключ-значение
        for key, value in request.POST.items():
             # Проверяем, является ли ключ строкой, представляющей положительное целое число
             if str(key).isdigit():
                 prices_from_request[key] = value # Сохраняем как {model_id_str: price_str}

        logger.debug(f"Prices received for update: {prices_from_request}")
        messages_list = [] # Собираем сообщения для AJAX ответа

        if not prices_from_request:
            messages_list.append({'tags': 'info', 'message': "Нет цен для обновления."})
            return render_and_return_partial(messages_list)

        # Получаем модели из БД, которые есть в запросе и в сессии, одним запросом
        # Фильтруем ключи из запроса, которые могут быть приведены к int
        model_ids_from_request = [int(mid) for mid in prices_from_request.keys() if str(mid).isdigit()]
        # Также добавим ID из текущей сессии, чтобы убедиться, что все выбранные модели загружены
        model_ids_in_session = [int(mid_s) for mid_s in selection.keys() if str(mid_s).isdigit()]
        # Объединяем и оставляем только уникальные ID
        all_relevant_model_ids = list(set(model_ids_from_request + model_ids_in_session))

        valid_models = ProductModel.objects.filter(id__in=all_relevant_model_ids).in_bulk() # Получаем словарь {id: model}
        logger.debug(f"Loaded {len(valid_models)} models relevant to update request/session.")


        is_session_modified = False # Флаг для отслеживания изменений в сессии

        # Итерируем по парам model_id_str, price_str, которые пришли в запросе
        # Важно: Итерируем по prices_from_request.items() чтобы знать, какие цены были отправлены.
        for model_id_str, price_str_raw in prices_from_request.items():
            # Только если model_id_str является цифрой и есть в текущей сессии
            # Проверка str(model_id_str).isdigit() уже сделана при формировании prices_from_request
            if model_id_str in selection:
                 model_id_int = int(model_id_str)
                 model = valid_models.get(model_id_int) # Получаем модель из предзагруженного словаря

                 if model: # Убедимся, что модель существует в БД (она должна, т.к. мы фильтровали по ID из сессии)
                     item_data = selection[model_id_str] # Получаем ссылку на данные из сессии

                     # Проверяем, что item_data - это словарь
                     if not isinstance(item_data, dict):
                          logger.warning(f"Skipping price update for item ID {model_id_str}: Invalid item_data format in session. Data: {item_data}")
                          messages_list.append({'tags': 'warning', 'message': f"Предупреждение: Некорректные данные для позиции (ID: {model_id_str})."})
                          # Возможно, стоит удалить некорректный элемент здесь? get_selected_items_and_total сделает это позже.
                          continue # Пропускаем этот элемент

                     price_str = price_str_raw.strip() # Обрезаем пробелы из введенной цены

                     try:
                         # Получаем текущую цену в сессии для сравнения
                         current_price_in_session_str = item_data.get('price')
                         # Если в сессии нет цены, используем цену из БД для сравнения
                         current_price_in_effect = Decimal(current_price_in_session_str) if current_price_in_session_str is not None and current_price_in_session_str != '' else model.price

                         if price_str == '': # Если цена пустая строка
                              # Если ключ 'price' существует в сессии (т.е. там была установлена цена), удаляем его
                              if 'price' in item_data:
                                   del item_data['price'] # Удаляем ключ, чтобы при расчете использовалась цена из БД
                                   # Нет необходимости переприсваивать item_data, т.к. item_data - это ссылка
                                   is_session_modified = True
                                   logger.debug(f"Price cleared for model ID {model_id_str}. Using DB price.")
                              else:
                                   logger.debug(f"Price was already empty/missing for model ID {model_id_str}. No change needed.")
                         else: # Если цена не пустая строка, пытаемся преобразовать в Decimal
                             price_value = Decimal(price_str)
                             if price_value >= 0: # Валидное число >= 0
                                  # Проверяем, изменилась ли цена по сравнению с тем, что БЫЛО в сессии (или из БД, если в сессии не было)
                                  # Сравниваем с Decimal значением
                                  if price_value != current_price_in_effect:
                                       # Обновляем цену в сессии
                                       item_data['price'] = str(price_value) # Сохраняем как строку
                                       # Нет необходимости переприсваивать item_data, т.к. item_data - это ссылка
                                       is_session_modified = True
                                       logger.debug(f"Price updated for model ID {model_id_str} to {price_value}. New item_data: {item_data}")
                                  else:
                                       logger.debug(f"Price for model ID {model_id_str} unchanged ({price_value}).")

                             else: # Негативная цена
                                  logger.warning(f"Received negative price '{price_str_raw}' for model ID {model_id_str}. Skipping update.")
                                  messages_list.append({'tags': 'warning', 'message': mark_safe(f"Отклонена отрицательная цена для позиции '<strong>{model.name}</strong>'. Использована текущая цена.")})

                     except (ValueError, InvalidOperation):
                         logger.warning(f"Invalid price format '{price_str_raw}' for model ID {model_id_str}. Skipping update.")
                         messages_list.append({'tags': 'warning', 'message': mark_safe(f"Некорректный формат цены для позиции '<strong>{model.name}</strong>'. Использована текущая цена.")})
                     except Exception as e:
                          logger.error(f"An unexpected error occurred while updating price for model ID {model_id_str}: {e}", exc_info=True)
                          messages_list.append({'tags': 'error', 'message': mark_safe(f"Неожиданная ошибка при обновлении цены для позиции '<strong>{model.name}</strong>'. Использована текущая цена.")})
                          continue # Продолжаем обработку других элементов
                 else:
                     # Модель не найдена в БД, хотя была в сессии (очень редкий случай)
                      if model_id_str in selection: # Если она была в сессии, удаляем на всякий случай
                          del selection[model_id_str]
                          is_session_modified = True
                          logger.warning(f"Model ID {model_id_str} not found in DB during apply_prices. Removed from session.")
                          messages_list.append({'tags': 'warning', 'message': f"Позиция с ID {model_id_str} не найдена и удалена из выбора."})
                      # else: модель пришла в запросе, но ее нет в сессии (нормально, просто игнорируем)

            # Если model_id_str не был в selection (не выбран), игнорируем его в apply_prices.
            # Мы обновляем цены только для УЖЕ выбранных позиций.

        # Сохраняем измененную сессию, если были изменения
        if is_session_modified:
             # selection словарь был изменен напрямую, т.к. item_data были ссылками
             request.session.modified = True
             logger.info("Session marked modified and saved after apply_prices.")
             logger.debug(f"Final selection state in session after apply_prices: {request.session.get('selection')}")
        else:
             logger.info("No session changes required during apply_prices.")


        # Возвращаем обновленную сводку. Сообщения будут добавлены в список messages_list
        # и отображены в частичном шаблоне.
        # messages.success(request, "Цены успешно применены." if not messages_list else "Цены применены с замечаниями.") # Стандартное сообщение
        # Добавляем общее сообщение в список ajax_messages
        if not messages_list:
             messages_list.append({'tags': 'success', 'message': "Цены успешно применены."})
        else:
             # Если были предупреждения/ошибки, добавляем сообщение, что применены с замечаниями
             messages_list.append({'tags': 'warning', 'message': "Цены применены с замечаниями."})


        return render_and_return_partial(messages_list) # Передаем список сообщений из цикла обработки


    # === Обработка действий, связанных с количеством ('add', 'remove', 'set', 'remove_all') ===
    # Для этих действий нужен один model_id_str
    model_id_str = request.POST.get('model_id')
    if not model_id_str or not str(model_id_str).isdigit():
         logger.error(f"Error: Missing or invalid model_id for action '{action}' in POST request.")
         return HttpResponseBadRequest(f"Отсутствует или некорректный ID модели для действия '{action}'.")

    try:
        model_id = int(model_id_str)
        # Получаем модель из БД. Она нужна для цены по умолчанию при добавлении/установке кол-ва,
        # а также для проверки существования.
        product_model = ProductModel.objects.select_related('product').get(id=model_id)
        logger.debug(f"Successfully retrieved ProductModel ID {model_id} ({product_model.name}) for action '{action}'.")

    except (ValueError, ProductModel.DoesNotExist) as e:
        error_message = f"Модель с ID {model_id_str} не найдена или некорректна: {e}"
        logger.error(f"Error processing model ID {model_id_str} for action '{action}': {e}")
        # Если модель не найдена, но она была в сессии (старый/некорректный ID), удаляем ее
        if model_id_str in selection:
             del selection[model_id_str]
             request.session['selection'] = selection
             request.session.modified = True
             logger.warning(f"Removed invalid/non-existent model ID {model_id_str} from session.")
             messages_list = [{'tags': 'warning', 'message': f"Предупреждение: Позиция с ID {model_id_str} не найдена и удалена из выбора."}]
        else:
             messages_list = [] # Нет сообщений, если элемента не было в сессии

        # Возвращаем частичную сводку после возможного удаления некорректного элемента
        return render_and_return_partial(messages_list)


    except Exception as e:
         error_message = f"Неожиданная ошибка при получении модели ID {model_id_str} для действия '{action}': {e}"
         logger.error(error_message, exc_info=True)
         return HttpResponse(error_message, status=500)


    # === Логика обновления сессии с новой структурой {'quantity': ..., 'price': ...} ===
    # Получаем текущие данные для модели из сессии или инициализируем, если нет.
    # При инициализации используем quantity 0 и цену из БД.
    item_data = selection.get(model_id_str, {'quantity': 0, 'price': str(product_model.price)})

    # Убедимся, что quantity в item_data - это число, иначе сбросим
    try:
        item_data['quantity'] = int(item_data.get('quantity', 0))
    except (ValueError, TypeError):
         item_data['quantity'] = 0
         logger.warning(f"Non-integer quantity found in session for model ID {model_id_str}. Resetting to 0.")
         # Сообщения об ошибках количества будут собраны в get_selected_items_and_total при следующем рендере
         # messages_list = [{'tags': 'warning', 'message': f"Некорректное количество для позиции '{product_model.name}'. Установлено 0."}]


    is_modified = False # Флаг для отслеживания изменений в item_data или selection
    messages_list = [] # Сообщения, специфичные для этих действий (сейчас нет, но может понадобиться)


    if action == 'add':
        item_data['quantity'] += 1
        is_modified = True
        # Если это первое добавление (было quantity 0 или не было в selection),
        # убедимся, что price инициализирован из БД (уже делается в .get, но на всякий случай)
        if 'price' not in item_data or item_data['price'] is None:
             item_data['price'] = str(product_model.price) # Убедимся, что цена задана (из БД)
             # is_modified = True # Уже True выше

        logger.debug(f"Action 'add': Added 1. New data: {item_data}.")


    elif action == 'remove':
        if item_data['quantity'] > 0:
            item_data['quantity'] -= 1
            is_modified = True
            logger.debug(f"Action 'remove': Removed 1. New data: {item_data}.")
        else:
             logger.debug("Action 'remove': Model not in selection or quantity already 0. No change.")

    elif action == 'remove_all':
        if model_id_str in selection:
            del selection[model_id_str]
            is_modified = True
            logger.debug(f"Action 'remove_all': Removed model ID {model_id_str} ({product_model.name}) completely.")
        else:
             logger.debug("Action 'remove_all': Model not in selection. No change.")

    elif action == 'set':
        quantity_str = request.POST.get('quantity')
        if quantity_str is None:
             logger.error("Action 'set': Missing quantity in POST data.")
             return HttpResponseBadRequest("Отсутствует значение количества для действия 'set'.")
        try:
            quantity = int(quantity_str)
            if quantity >= 0:
                if quantity > 0:
                    # Обновляем item_data только если количество > 0
                    item_data['quantity'] = quantity
                    is_modified = True
                    # Если устанавливаем количество (>0) и цена еще не задана/инициализирована, берем из БД
                    if 'price' not in item_data or item_data['price'] is None:
                         item_data['price'] = str(product_model.price) # Убедимся, что цена задана (из БД)
                         # is_modified = True # Уже True выше
                    logger.debug(f"Action 'set': Quantity set to {quantity}. Data: {item_data}.")
                elif quantity == 0:
                     # Если устанавливаем 0, удаляем элемент из selection, если он там есть
                     if model_id_str in selection:
                        del selection[model_id_str]
                        is_modified = True
                        logger.debug("Action 'set': Quantity set to 0, removed from selection.")
                     else:
                         logger.debug("Action 'set': Quantity set to 0 for model not in selection. No change.")
            else:
                logger.warning(f"Action 'set': Invalid quantity (negative): {quantity_str}.")
                return HttpResponseBadRequest("Количество должно быть неотрицательным числом.")

        except ValueError:
            logger.warning(f"Action 'set': Invalid quantity format: '{quantity_str}'.")
            return HttpResponseBadRequest(f"Некорректное значение количества: '{quantity_str}'. Должно быть целое число.")

    else:
        # Этот блок должен быть достигнут только если action не 'clear', не 'apply_prices'
        # и не одно из действий количества.
        logger.error(f"Error: Unknown action '{action}' with model_id '{model_id_str}'.")
        return HttpResponseBadRequest(f"Неизвестное действие: '{action}'")

    # Если item_data был изменен и модель все еще должна быть в selection (quantity > 0)
    # При действиях 'add', 'remove', 'set' > 0
    # При действиях 'remove_all' или 'set' == 0, мы уже удалили из selection выше
    if is_modified and action in ['add', 'remove', 'set'] and item_data['quantity'] > 0:
         selection[model_id_str] = item_data # Сохраняем обновленный item_data в selection
         logger.debug(f"Item {model_id_str} saved back to selection with updated data.")


    # Сохраняем измененную сессию только если были внесены изменения
    if is_modified:
        request.session['selection'] = selection # Сохраняем обновленный словарь selection
        request.session.modified = True
        logger.debug(f"Session marked modified and saved after action '{action}'. Final state: {request.session.get('selection')}")
    else:
        logger.debug(f"No session changes made for action '{action}'.")

    # === После успешного обновления, возвращаем обновленный HTML сводки ===
    return render_and_return_partial(messages_list) # Передаем любые собранные сообщения


# --- Хелпер функции для PDF ---
# (Без изменений относительно предыдущего полного файла)
def get_font_base64(font_path):
    """Читает файл шрифта по локальному пути и возвращает его содержимое в Base64."""
    try:
        with open(font_path, 'rb') as font_file:
            encoded_string = base64.b64encode(font_file.read()).decode('ascii')
            return encoded_string
    except FileNotFoundError:
        logger.error(f"Error: Font file NOT FOUND at {font_path}")
        return None
    except Exception as e:
        logger.error(f"Error encoding font file {font_path}: {e}", exc_info=True)
        return None

def get_image_base64(image_path):
    """Reads an image file by its local path and returns its Base64 Data URI (for PDF)."""
    global PIL_INSTALLED
    try:
        with open(image_path, 'rb') as image_file:
            if PIL_INSTALLED:
                try:
                    img = Image.open(image_file)
                    mime_type = Image.MIME.get(img.format, 'image/png')
                    image_file.seek(0) # Вернуть указатель файла в начало после определения формата
                except Exception as pil_e:
                     logger.warning(f"Warning: Pillow failed to identify image format for {image_path}: {pil_e}. Falling back to extension.", exc_info=True)
                     PIL_INSTALLED = False # Отключаем Pillow, если произошла ошибка
                finally:
                     # Если Pillow не сработало или не установлено
                     if not PIL_INSTALLED or 'mime_type' not in locals() or mime_type is None:
                        extension = os.path.splitext(image_path)[1].lower()
                        if extension == '.png': mime_type = 'image/png'
                        elif extension in ['.jpg', '.jpeg']: mime_type = 'image/jpeg'
                        elif extension == '.gif': mime_type = 'image/gif'
                        else:
                             logger.warning(f"Warning: Unknown image extension '{extension}' for {image_path} and MIME type detection failed. Defaulting to image/png.")
                             mime_type = 'image/png'
            else: # Если Pillow не установлен изначально
                extension = os.path.splitext(image_path)[1].lower()
                if extension == '.png': mime_type = 'image/png'
                elif extension in ['.jpg', '.jpeg']: mime_type = 'image/jpeg'
                elif extension == '.gif': mime_type = 'image/gif'
                else:
                     logger.warning(f"Warning: Unknown image extension '{extension}' for {image_path} and Pillow not installed. Defaulting to image/png.")
                     mime_type = 'image/png'

            encoded_string = base64.b64encode(image_file.read()).decode('ascii')
            data_uri = f'data:{mime_type};base64,{encoded_string}'
            return data_uri
    except FileNotFoundError:
        logger.error(f"Error: Image file NOT FOUND at {image_path}")
        return None
    except Exception as e:
        logger.error(f"Error encoding image file {image_path}: {e}", exc_info=True)
        return None


# --- Представление для генерации PDF (использует get_selected_items_and_total) ---
@login_required
def generate_commercial_offer_pdf_view(request):
    """
    Генерирует PDF коммерческого предложения с выбранными моделями,
    их описаниями (details), изображениями и ценами из сессии.
    """
    logger.info("\nEntering generate_commercial_offer_pdf_view...")

    selection = request.session.get('selection', {})
    # Используем общую функцию для сбора данных и расчета сумм
    # Передаем копию сессии, чтобы она не изменилась, если будут обнаружены некорректные элементы
    selected_items_data, total_price, messages_from_processing, _ = get_selected_items_and_total(selection.copy()) # Не используем updated_selection_copy


    # Добавляем сообщения из обработки в стандартные Django messages
    for msg_data in messages_from_processing:
         level = messages.WARNING if msg_data.get('tags') == 'warning' else messages.ERROR
         messages.add_message(request, level, msg_data['message'])


    if not selected_items_data:
        messages.info(request, "Ваш выбор пуст. Невозможно сформировать PDF.")
        logger.warning("Selection is empty for PDF generation, redirecting.")
        return redirect('catalog_selection')

    # === Логируем создание документа ===
    document_number = None
    try:
        document_number = get_next_document_number()
        log_entry = DocumentLog(user=request.user, document_number=document_number, document_type='pdf')
        log_entry.save()
        logger.info(f"Logged PDF generation: Number={document_number}, User={request.user.username}")
    except Exception as e:
        logger.error(f"Error logging PDF generation: {e}", exc_info=True)
        messages.error(request, f"Ошибка при записи в журнал документов: {e}. Документ может быть сгенегерирован без номера.")

    # === Определяем строковые представления даты и номера документа ===
    doc_date_str = datetime.date.today().strftime('%d.%m.%Y')
    doc_number_str = str(document_number) if document_number is not None else ""

    # --- Получение локальных путей к файлам шрифтов DejaVu Sans ---
    font_dir_path = settings.BASE_DIR / 'static' / 'fonts'
    logger.debug(f"Checking font directory: {font_dir_path}")

    dejavu_sans_regular_path = font_dir_path / 'DejaVuSans.ttf'
    dejavu_sans_bold_path = font_dir_path / 'DejaVuSans-Bold.ttf'

    dejavu_sans_regular_base64 = get_font_base64(str(dejavu_sans_regular_path))
    dejavu_sans_bold_base64 = get_font_base64(str(dejavu_sans_bold_path))

    if not dejavu_sans_regular_base64 or not dejavu_sans_bold_base64:
        error_msg = "Не удалось загрузить необходимые шрифты DejaVu Sans для PDF. Проверьте файлы шрифтов в static/fonts."
        logger.error(error_msg)
        messages.error(request, error_msg)
        return redirect('catalog_selection')

    # Добавляем данные изображений (base64) к selected_items_data
    # Делаем это здесь, чтобы не вызывать get_image_base64 в get_selected_items_and_total
    # (т.к. там не всегда нужно изображение, а здесь нужно только для PDF)
    for item in selected_items_data:
         image_data_uri = None
         model = item['model']
         # В PDF мы используем Base64, а не локальный путь
         if model.image and hasattr(model.image, 'path') and os.path.exists(model.image.path):
              image_path = model.image.path # Получаем локальный путь для чтения
              image_data_uri = get_image_base64(image_path)
         elif model.image:
              logger.warning(f"Warning: Image file for Model ID {model.id} ('{model.name}') ({model.image.name}) not found on disk at {getattr(model.image, 'path', 'N/A')}")

         item['image_data_uri'] = image_data_uri # Добавляем ключ image_data_uri в словарь элемента

    context = {
        'selected_items': selected_items_data, 'total_price': total_price, 'offer_title': 'Коммерческое предложение',
        'delivery_terms': '20 рабочих дней', 'warranty_terms': '12 месяцев',
        'manager_name': 'Логинов Алексей', 'company_name': 'ЛУЧ-IT (ООО «Луч АйТи)',
        'manager_contact': '+7-382-299-56-13', 'manager_email': 'komdir@luch-it.ru',
        'manager_signature_name': 'Логинов А. А.', 'attorney_details': 'Доверенность от 01.02.2024г',
        'document_date': datetime.date.today(), 'document_number': document_number,
        'dejavu_sans_regular_base64': dejavu_sans_regular_base64,
        'dejavu_sans_bold_base64': dejavu_sans_bold_base64,
    }
    logger.info("Context prepared for PDF template.")

    html_string = render_to_string('catalog/commercial_offer_pdf_template.html', context)

    result = BytesIO()
    logger.info("Calling pisa.CreatePDF...")
    try:
        pisa_status = pisa.CreatePDF(
           html_string.encode('utf-8'), dest=result, encoding='utf-8',
        )
        logger.info("pisa.CreatePDF call finished.")
    except Exception as e:
        logger.error(f"\n{'='*50}\nPISA CreatePDF threw an exception: {e}\n{'='*50}\n", exc_info=True)
        messages.error(request, f"Критическая ошибка при вызове pisa.CreatePDF: {e}. Проверьте логи сервера.")
        return redirect('catalog_selection')

    if pisa_status.err:
        logger.error("\n" + "="*50)
        logger.error("PISA Error during PDF generation (pisa_status.err is True):")
        # pisa_status.err может быть списком ошибок или одиночным объектом
        if isinstance(pisa_status.err, list):
             for i, err in enumerate(pisa_status.err):
                 logger.error(f"Error {i+1}: {err}")
                 # Некоторые объекты ошибок могут иметь атрибуты msg или url
                 if hasattr(err, 'msg'): logger.error(f"  Message: {err.msg}")
                 if hasattr(err, 'url'): logger.error(f"  URL: {err.url}")
        else: # Одиночная ошибка
             logger.error(f"Single Error: {pisa_status.err}")
             if hasattr(pisa_status.err, 'msg'): logger.error(f"  Message: {pisa_status.err.msg}")
             if hasattr(pisa_status.err, 'url'): logger.error(f"  URL: {pisa_status.err.url}")
        logger.error("="*50 + "\n")
        messages.error(request, "Произошла ошибка при генерации PDF файла. Проверьте логи сервера.")
        return redirect('catalog_selection')

    pdf_bytes = result.getvalue()
    response = HttpResponse(pdf_bytes, content_type='application/pdf')
    filename = f"commercial_offer_{doc_number_str or 'undated'}.pdf"
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    logger.info("PDF generated successfully. Returning HttpResponse.")
    return response


# === Хелпер функция для добавления границ к ячейке DOCX ---
# (Без изменений относительно предыдущего полного файла)
def set_cell_border(cell, **kwargs):
    """Применяет границы к ячейке Word."""
    # Ensure necessary imports are available globally
    # from docx.oxml.shared import OxmlElement, qn # Imported at top
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    borders = {
        'top': {'val': 'single', 'sz': 12, 'color': '#000000'},
        'bottom': {'val': 'single', 'sz': 12, 'color': '#000000'},
        'left': {'val': 'single', 'sz': 12, 'color': '#000000'},
        'right': {'val': 'single', 'sz': 12, 'color': '#000000'},
    }

    # Update default borders with provided kwargs
    for border_name, border_props in kwargs.items():
        if border_name in borders:
             borders[border_name] = border_props
        elif border_name in ['top', 'bottom', 'left', 'right'] and border_props is None:
             borders[border_name] = None

    tcBorders = tcPr.first_child_found_in('w:tcBorders')
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    for border_name, border_props in borders.items():
        border_element = tcBorders.find(qn('w:' + border_name))
        if border_props is not None:
            if border_element is None:
                border_element = OxmlElement('w:' + border_name)
                tcBorders.append(border_element)
            # Set attributes from the properties dictionary
            for key, value in border_props.items():
                 border_element.set(qn('w:' + key), str(value))
        else:
            # If border_props is None, remove the border element
            if border_element is not None:
                 tcBorders.remove(border_element)


# --- Универсальная функция для генерации DOCX ---
def generate_docx_offer(request, template_file_name, filename_prefix):
    """
    Генерирует DOCX коммерческого предложения по указанному шаблону,
    используя цены из сессии и старую логику загрузки изображений.
    """
    print(f"\nEntering generate_docx_offer with template: {template_file_name}")

    # --- Сбор данных о выбранных товарах (СТАРАЯ ЛОГИКА ДЛЯ ИЗОБРАЖЕНИЙ) ---
    selection = request.session.get('selection', {})
    selected_items_data = []
    # Используем Decimal для точных расчетов цены
    total_price = Decimal('0.00')

    if not selection:
        messages.info(request, "Ваш выбор пуст. Невозможно сформировать DOCX.")
        print("Selection is empty, redirecting.")
        return redirect('catalog_selection')

    print("Starting to collect selected items data (using old image logic)...")
    # Получаем все выбранные модели одним запросом для эффективности
    model_ids_in_selection = [int(model_id_str) for model_id_str in selection.keys() if str(model_id_str).isdigit()]
    selected_models_queryset = ProductModel.objects.select_related('product').filter(id__in=model_ids_in_selection)
    selected_models_dict = {str(model.id): model for model in selected_models_queryset}


    for model_id_str, item_data in list(selection.items()): # Итерируем по элементам сессии
         try:
            model_id = int(model_id_str)
            model = selected_models_dict.get(model_id_str) # Получаем модель из предзагруженного словаря

            if not model:
                 print(f"Error: Model ID {model_id_str} not found in DB during DOCX generation. Removing from selection.")
                 if model_id_str in selection:
                     del selection[model_id_str]
                     request.session.modified = True
                 messages.warning(request, f"Предупреждение: Позиция с ID {model_id_str} не найдена и удалена из выбора.")
                 continue

            # Читаем quantity и price из item_data словаря
            quantity = item_data.get('quantity', 0)
            price_str = item_data.get('price')

            try:
                 quantity = int(quantity) if quantity is not None and str(quantity).strip().isdigit() else 0
                 if quantity < 0: quantity = 0

                 # Используем цену из сессии, если она есть и корректна, иначе из БД
                 unit_price = model.price # Цена по умолчанию из БД
                 if price_str is not None and price_str != '':
                      try:
                           price_from_session = Decimal(price_str)
                           if price_from_session >= 0:
                                unit_price = price_from_session # Используем цену из сессии, если валидна
                           else:
                                logger.warning(f"Negative price ({price_from_session}) found in session for model ID {model_id_str}. Using DB price.")
                                messages.warning(request, f"Обнаружена некорректная цена для позиции '{model.name}'. Использована цена из каталога.")
                      except (ValueError, InvalidOperation):
                           logger.warning(f"Invalid price format '{price_str}' in session for model ID {model_id_str}. Using DB price.")
                           messages.warning(request, f"Предупреждение: некорректная цена для позиции '{model.name}' в выборе. Использована цена из каталога.")
                 # Если price_str is None or '', unit_price остается из БД

            except (ValueError, InvalidOperation):
                 logger.warning(f"Error converting quantity ({quantity}) or price ({price_str}) for model ID {model_id_str}. Skipping item.")
                 if model_id_str in selection: # Удаляем на всякий случай
                     del selection[model_id_str]
                     request.session.modified = True
                 continue # Пропускаем элемент

            if quantity <= 0:
                 print(f"Warning: Found zero or negative quantity ({quantity}) for model ID {model_id_str} during DOCX generation. Removing from selection.")
                 if model_id_str in selection:
                     del selection[model_id_str]
                     request.session.modified = True
                 continue

            # === СТАРАЯ ЛОГИКА ПОЛУЧЕНИЯ IMAGE_PATH ===
            image_path = None
            # Проверяем, есть ли image, атрибут .path и СУЩЕСТВУЕТ ли файл на диске
            if model.image and hasattr(model.image, 'path') and os.path.exists(model.image.path):
                 image_path = model.image.path # <-- Если все ОК, ПУТЬ ПОЛУЧАЕТСЯ ЗДЕСЬ
                 print(f"DEBUG: Image path found on disk: {image_path}")
            elif model.image:
                 # Если image есть, но .path нет ИЛИ os.path.exists(model.image.path) == False
                 # Выводится предупреждение БЕЗ ИСПОЛЬЗОВАНИЯ image_path локально в этом блоке
                 print(f"Warning: Image file for Model ID {model.id} ({model.image.name}) not found on disk at {getattr(model.image, 'path', 'N/A')}")
            else:
                 print(f"DEBUG: Model {model.name} has no image attached.")
            # =========================================


            item_total = unit_price * quantity
            # === ДОБАВЛЕНО ДЛЯ ОТЛАДКИ ===
            print(f"DEBUG: Item {model.name} (ID: {model.id}) - Effective Price: {unit_price}, Quantity: {quantity}, Calculated item_total: {item_total}")
            # ==============================

            selected_items_data.append({
                'model': model,
                'quantity': quantity,
                'unit_price': unit_price, # Используем эффективную цену
                'item_total': item_total,
                'image_path': image_path, # <-- image_path (может быть None) ДОБАВЛЯЕТСЯ в словарь item
            })
            total_price += item_total

         except (ProductModel.DoesNotExist, ValueError) as e:
             if model_id_str in selection:
                 print(f"Error processing model ID {model_id_str} for DOCX: {e}. Removing from selection.")
                 del selection[model_id_str]
                 request.session.modified = True
                 messages.warning(request, f"Одна из выбранных позиций (ID: {model_id_str}) не найдена или некорректна и была удалена.")
             continue
         except Exception as e:
              print(f"An unexpected error occurred while processing model ID {model_id_str} for DOCX: {e}")
              messages.error(request, f"Произошла неожиданная ошибка при обработке позиции (ID: {model_id_str}): {e}")
              continue

    request.session.modified = True # Убеждаемся, что изменения в сессии (удаление некорректных) сохранятся

    if not selected_items_data:
         messages.warning(request, "В вашем выборе нет корректных позиций для формирования DOCX.")
         print("selected_items_data is empty after processing, redirecting.")
         return redirect('catalog_selection')

    # === Логируем создание документа ===
    document_number = None
    try:
        document_number = get_next_document_number()
        log_entry = DocumentLog(
            user=request.user,
            document_number=document_number,
            document_type='docx'
        )
        log_entry.save()
        print(f"Logged DOCX document generation: Number={document_number}, User={request.user.username}, Template={template_file_name}")
    except Exception as e:
        print(f"Error logging DOCX generation: {e}")
        messages.error(request, f"Ошибка при записи в журнал документов: {e}. Документ может быть сгенерирован без номера.")

    # === Определяем строковые представления даты и номера документа ===
    doc_date_str = datetime.date.today().strftime('%d.%m.%Y')
    doc_number_str = str(document_number) if document_number is not None else ""

    # === Определяем путь к шаблонному файлу DOCX ===
    template_path = settings.BASE_DIR / 'templates' / template_file_name
    print(f"Attempting to load DOCX template from: {template_path}")

    if not template_path.exists():
        error_msg = f"Критическая ошибка: Шаблонный файл DOCX не найден по пути {template_path}. Невозможно сгенерировать документ без шаблона."
        print(error_msg)
        messages.error(request, error_msg + f" Убедитесь, что файл '{template_file_name}' существует в папке templates вашего проекта.")
        return redirect('catalog_selection')

    try:
        document = Document(template_path)
        print(f"Successfully loaded DOCX template from {template_path}")
    except Exception as e:
        error_msg = f"Ошибка при загрузке шаблонного DOCX файла из {template_path}: {e}"
        print(error_msg)
        messages.error(request, error_msg)
        return redirect('catalog_selection')

    # --- Замена текстовых плейсхолдеров ---
    placeholders = {
        '{{document_date}}': doc_date_str,
        '{{document_number}}': doc_number_str if doc_number_str else 'б/н',
        # TODO: Добавьте плейсхолдеры для других полей, если они есть в шаблоне DOCX
        # '{{offer_title}}': 'Коммерческое предложение',
        # '{{delivery_terms}}': '20 рабочих дней',
        # '{{warranty_terms}}': '12 месяцев',
        # '{{manager_name}}': 'Логинов Алексей',
        # '{{company_name}}': 'ЛУЧ-IT (ООО «Луч АйТи)',
        # '{{manager_contact}}': '+7-382-299-56-13',
        # '{{manager_email}}': 'komdir@luch-it.ru',
        # '{{manager_signature_name}}': 'Логинов А. А.',
        # '{{attorney_details}}': 'Доверенность от 01.02.2024г',
    }

    def simple_replace_placeholders_in_paragraphs(paragraphs_list):
        for paragraph in paragraphs_list:
            paragraph_text = paragraph.text
            replaced_text = paragraph_text
            for key, value in placeholders.items():
                if key in replaced_text:
                    replaced_text = replaced_text.replace(key, str(value))
            if replaced_text != paragraph_text:
                 paragraph.clear()
                 paragraph.add_run(replaced_text)

    simple_replace_placeholders_in_paragraphs(document.paragraphs)
    for section in document.sections:
         if section.header: simple_replace_placeholders_in_paragraphs(section.header.paragraphs)
         if section.footer: simple_replace_placeholders_in_paragraphs(section.footer.paragraphs)
    for tbl in document.tables:
        for row in tbl.rows:
            for cell in row.cells:
                simple_replace_placeholders_in_paragraphs(cell.paragraphs)
    print("Placeholder replacement finished.")


    # --- Вставка таблицы товаров (ищем плейсхолдер) ---
    table_placeholder_text = '[PRODUCT_TABLE]'
    table_placeholder_paragraph = None
    for p in document.paragraphs:
        if table_placeholder_text in p.text:
            table_placeholder_paragraph = p
            break

    if not table_placeholder_paragraph:
        error_msg = f"Ошибка: В шаблонном файле DOCX '{template_file_name}' не найден маркер '{table_placeholder_text}'. Убедитесь, что он присутствует в шаблоне в основном теле документа."
        print(error_msg)
        messages.error(request, error_msg)
        print("Skipping table insertion because placeholder was not found.")
        output = BytesIO()
        try:
            document.save(output)
            output.seek(0)
            response = HttpResponse(output.getvalue(),
                                    content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            filename = f"commercial_offer_{filename_prefix}_{doc_number_str or 'undated'}_no_table.docx"
            response['Content-Disposition'] = f'attachment; filename="{filename}"'
            print("DOCX generated successfully (No table - placeholder missing). Returning HttpResponse.")
            return response
        except Exception as e:
            save_error_msg = f"Ошибка при сохранении DOCX файла без таблицы ('{template_file_name}'): {e}."
            print(save_error_msg)
            messages.error(request, save_error_msg)
            return HttpResponse("Error saving DOCX file without table: " + str(e), status=500)


    # === Определяем точку вставки таблицы и удаляем плейсхолдер ===
    parent_element = table_placeholder_paragraph._element.getparent()
    placeholder_index = parent_element.index(table_placeholder_paragraph._element)
    parent_element.remove(table_placeholder_paragraph._element)


    # === Создаем НОВУЮ таблицу и вставляем ее на место плейсхолдера ===
    print("Creating table with 1 row for headers...")
    num_cols = 4
    table = document.add_table(rows=1, cols=num_cols) # 4 колонки: Товар, Кол-во, Цена, Сумма

    # Вставляем таблицу на место удаленного абзаца-плейсхолдера
    table_element = table._element
    body = document._body._element
    try:
        parent_element.insert(placeholder_index, table_element)
        print(f"Inserted table element at index {placeholder_index}.")
    except Exception as insert_e:
        print(f"Error inserting table element at placeholder index {placeholder_index}: {insert_e}. It might remain at the end.")
        # Если вставка не удалась, таблица, вероятно, осталась в конце


    # Устанавливаем примерные ширины колонок (опционально)
    # Ширины в CM. Должно быть ровно столько, сколько колонок.
    col_widths_cm = [Cm(7), Cm(2), Cm(3), Cm(3)] # Примерные ширины для 4 колонок
    # В старой версии не было проверки WD_TABLE_COL_WIDTH и len(col_widths_cm) == num_cols
    # Просто пытались установить ширину, что могло вызвать ошибку если не 4 колонки или нет WD_TABLE_COL_WIDTH
    # Оставим как было в предоставленной старой версии
    try:
        for col_idx, width in enumerate(col_widths_cm):
             if col_idx < len(table.columns): # Проверяем, что колонка существует
                 table.columns[col_idx].width = width
                 # table.columns[col_idx].preferred_width_type = WD_TABLE_COL_WIDTH.W # В старой версии этого не было
                 # table.columns[col_idx].preferred_width = width.emu # В старой версии этого не было
        # table.preferred_width_type = WD_TABLE_COL_WIDTH.PCT # В старой версии этого не было
        # table.preferred_width = 10000 # В старой версии этого не было
        print("Table column widths set (old logic).")
    except Exception as e:
        # В старой версии логирование было print
        print(f"Warning: Could not set table column widths programmatically (old logic): {e}")
        print("If borders are missing, ensure they are set in the template file.")


    # --- Заполнение таблицы ---
    # Заголовки таблицы (заполняем первую строку, созданную при init)
    print("Filling table headers...")
    hdr_cells = table.rows[0].cells
    # В старой версии не было проверки len(hdr_cells) >= num_cols здесь
    hdr_cells[0].text = 'Товар'
    hdr_cells[1].text = 'Кол-во, шт.'
    hdr_cells[2].text = 'Цена за ед., руб.'
    hdr_cells[3].text = 'Сумма, руб.'
    for cell in hdr_cells: # Здесь итерировались по всем ячейкам, даже если их больше 4
        if cell.paragraphs:
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                 run.bold = True
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    print("Table headers filled.")


    # Заполнение строк таблицы данными (используем add_row() для каждого элемента)
    print(f"Adding {len(selected_items_data)} rows for items...")
    for i, item in enumerate(selected_items_data):
       print(f"  Adding row for item {i+1}/{len(selected_items_data)}: {item['model'].name}") # В старой версии лог был проще
       # === Добавляем новую строку для каждого элемента ===
       # В старой версии не было проверки len(table.columns) >= num_cols здесь
       row_cells = table.add_row().cells # Каждая новая ячейка содержит один пустой абзац.

       # === Заполнение первой ячейки (Наименование / Описание / Изображение) ===
       product_info_cell = row_cells[0]
       # Исправленная логика: Очищаем все существующие абзацы, затем добавляем новые
       # Важно: Итерируем по копии списка, чтобы избежать проблем при удалении элементов во время итерации
       for p in list(product_info_cell.paragraphs):
            # В старой версии не было проверки p._element.getparent() is not None
            p._element.getparent().remove(p._element)

       # 1. Абзац для изображения (если есть)
       # Условие как в старой версии, используем image_path из item
       if item['image_path'] and os.path.exists(item['image_path']): # <-- УСЛОВИЕ КАК В СТАРОЙ ВЕРСИИ
           try:
               p_image = product_info_cell.add_paragraph()
               run = p_image.add_run()
               image_width_cm = 4.0
               run.add_picture(str(item['image_path']), width=Cm(image_width_cm)) # <-- ИСПОЛЬЗУЕМ image_path ИЗ ITEM
               p_image.alignment = WD_ALIGN_PARAGRAPH.CENTER
               p_image.paragraph_format.space_after = Cm(0.1)
               # В старой версии здесь не было отдельного log.debug о добавлении
           except Exception as e:
               # В старой версии логирование было print
               print(f"    Error adding image {item['image_path']} to DOCX table cell: {e}")
               img_name = getattr(item['model'].image, 'name', 'имя файла неизвестно') if item['model'].image else 'файл изображения не прикреплен'
               p_error = product_info_cell.add_paragraph()
               # ИСПРАВЛЕНО: apply italic = True directly to the run object returned by add_run
               p_error.add_run(f"[Ошибка загрузки изображения: {img_name}]").italic = True
               p_error.paragraph_format.space_after = Cm(0.1)
       # Условие как в старой версии
       elif item['model'].image: # <-- УСЛОВИЕ КАК В СТАРОЙ ВЕРСИИ
            img_name = getattr(item['model'].image, 'name', 'имя файла неизвестно')
            # В старой версии логирование было print и путь получался через getattr
            print(f"Warning: Image file for Model ID {item['model'].id} ({item['model'].image.name}) not found on disk at {getattr(item['model'].image, 'path', 'N/A')}")
            p_warning = product_info_cell.add_paragraph()
            # ИСПРАВЛЕНО: apply italic = True directly to the run object returned by add_run
            # В старой версии здесь была ваша ошибка .runs[-1]
            warning_run = p_warning.add_run(f"[Изображение модели '{item['model'].name}' не найдено на диске: {img_name}]")
            if warning_run: # Добавил проверку на всякий случай
                 warning_run.italic = True
            p_warning.paragraph_format.space_after = Cm(0.1)
       # В старой версии не было else для случая, когда image объекта нет вообще


       # 2. Абзац для Названия Модели
       print(f"  Adding name paragraph for {item['model'].name}...") # Добавил лог для проверки
       p_name = product_info_cell.add_paragraph()
       p_name.add_run(f"{item['model'].product.name} - {item['model'].name}").bold = True
       p_name.paragraph_format.space_after = Cm(0.1)
       print("    Name paragraph added.") # Добавил лог для проверки


       # 3. Абзацы для Описания Модели
       if item['model'].details:
           print(f"  Adding details paragraphs for {item['model'].name}...") # Добавил лог для проверки
           details_text = item['model'].details
           for j, line in enumerate(details_text.splitlines()):
                line = line.strip()
                if line:
                     p_details = product_info_cell.add_paragraph()
                     p_details.add_run(line).italic = True
                     p_details.paragraph_format.space_before = Cm(0.05)
                     p_details.paragraph_format.space_after = Cm(0.05)
           print(f"    Details paragraphs added ({details_text.count('\\n')+1} lines).") # Добавил лог для проверки
       else:
           print(f"  No details for {item['model'].name}. Skipping details paragraphs.") # Добавил лог для проверки


       # Вертикальное выравнивание для ячейки
       product_info_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

       # === Конец заполнения первой ячейки ===


       # Заполнение остальных ячеек в этой строке (Кол-во, Цена, Сумма)
       # Используем .text для установки текста, а затем форматируем первый абзац.
       # Это более простой и, возможно, более надежный способ для этих ячеек.

       qty_cell = row_cells[1]
       qty_cell.text = str(item['quantity'])
       if qty_cell.paragraphs:
            qty_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
       qty_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
       print(f"    Qty cell filled ({item['quantity']}).") # Добавил лог


       price_cell = row_cells[2]
       price_cell.text = f"{item['unit_price']:.2f}"
       if price_cell.paragraphs:
            price_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
       price_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
       print(f"    Price cell filled ({item['unit_price']:.2f}).") # Добавил лог


       total_cell = row_cells[3]
       total_cell.text = f"{item['item_total']:.2f}"
       if total_cell.paragraphs:
            # Делаем текст суммы жирным
            for run in total_cell.paragraphs[0].runs:
                 run.bold = True
            total_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
       total_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
       print(f"    Total cell filled ({item['item_total']:.2f}).") # Добавил лог


    # Строка Итого (добавляем новую строку)
    print("Adding total row...")
    # В старой версии не было проверки len(table.columns) >= num_cols здесь
    total_row = table.add_row()
    total_row_cells = total_row.cells

    # Объединяем ячейки (первые три)
    # В старой версии не было проверок на наличие ячеек перед объединением и try/except
    merged_total_cell = total_row_cells[0].merge(total_row_cells[1]).merge(total_row_cells[2])

    # === ИСПРАВЛЕННАЯ ЛОГИКА: ЗАПОЛНЕНИЕ ОБЪЕДИНЕННОЙ ЯЧЕЙКИ ДЛЯ "ИТОГО:" ===
    # Устанавливаем текст напрямую, это создаст/заменит первый абзац.
    merged_total_cell.text = "Итого:"
    # Затем форматируем первый абзац.
    if merged_total_cell.paragraphs:
        paragraph = merged_total_cell.paragraphs[0]
        for run in paragraph.runs: # Делаем весь текст в абзаце жирным
             run.bold = True
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT # Выравнивание по правому краю
    # Вертикальное выравнивание ячейки
    merged_total_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    print("Total label cell added.") # Добавил лог


    # Заполняем последнюю ячейку (4-ю) в строке Итого общей суммой
    final_total_cell = total_row_cells[3]

    # === ДОБАВЛЕНО ДЛЯ ОТЛАДКИ ===
    print(f"DEBUG: Calculated total_price before insertion into DOCX: {total_price}, type: {type(total_price)}")
    # ==============================

    # === ИСПРАВЛЕННАЯ ЛОГИКА: ЗАПОЛНЕНИЕ ЯЧЕЙКИ С СУММОЙ ИТОГО ===
    # Устанавливаем текст напрямую, это создаст/заменит первый абзац.
    final_total_cell.text = f"{total_price:.2f}"
    # Затем форматируем первый абзац.
    if final_total_cell.paragraphs:
        paragraph = final_total_cell.paragraphs[0]
        for run in paragraph.runs: # Делаем весь текст в абзаце жирным
            run.bold = True
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT # Выравнивание по правому краю
    # Вертикальное выравнивание ячейки
    final_total_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    print("Final total cell added.") # Добавил лог


    # === Применяем границы ко всем ячейкам таблицы программно ===
    # В старой версии не было try/except здесь
    print("Applying borders to table cells programmatically...")
    try: # Добавил try/except из новой версии
        # В старой версии не было проверки table.rows и row.cells
        for row in table.rows:
            for cell in row.cells:
                set_cell_border(cell)
        print("Programmatic borders applied.")
    except Exception as e:
        print(f"Error applying programmatic borders: {e}")
        messages.warning(request, f"Не удалось применить границы к таблице программно: {e}. Убедитесь, что библиотека python-docx установлена корректно и имеет необходимые компоненты.")
        print(f"If borders are missing, ensure they are set in the template file: {template_path}")
        messages.info(request, "Если границы таблицы не появились, настройте их в шаблонном файле DOCX через Microsoft Word.")


    # --- Сохранение и возврат HttpResponse ---
    print("Saving modified DOCX to BytesIO...")

    output = BytesIO()
    try:
        document.save(output)
        output.seek(0)
        print(f"Generated DOCX bytes length: {len(output.getvalue())} bytes.")
    except Exception as e:
         error_msg = f"Ошибка при сохранении DOCX файла ('{template_file_name}'): {e}."
         print(error_msg)
         messages.error(request, error_msg)
         return HttpResponse("Error saving DOCX file: " + str(e), status=500)

    filename = f"commercial_offer_{filename_prefix}_{doc_number_str or 'undated'}.docx"
    response = HttpResponse(output.getvalue(),
                            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{filename}"'

    print("DOCX generated successfully. Returning HttpResponse.")
    return response


# --- Представления-обертки для универсальной функции ---

@login_required
def generate_commercial_offer_docx_view(request):
    """Генерирует основной вариант DOCX КП."""
    return generate_docx_offer(request, 'commercial_offer_template.docx', 'main')

@login_required
def generate_commercial_offer_umed_docx_view(request):
    """Генерирует DOCX КП для Учебного мира."""
    return generate_docx_offer(request, 'commercial_offer_umed_template.docx', 'umed')

@login_required
def generate_commercial_offer_pos78_docx_view(request):
    """Генерирует DOCX КП для ПОС78."""
    return generate_docx_offer(request, 'commercial_offer_pos78_template.docx', 'pos78')


@user_passes_test(lambda u: u.is_staff)
@login_required
def document_log_view(request):
    """Отображает список всех созданных записей DocumentLog для администраторов."""
    logger.info(f"Executing document_log_view for user {request.user.username}")
    logs = DocumentLog.objects.all().order_by('-timestamp')
    context = {'logs': logs}
    logger.info(f"Found {logs.count()} document log entries. Rendering document_log.html...")
    return render(request, 'catalog/document_log.html', context)